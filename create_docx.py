#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_background(cell, fill_color):
    """Установить цвет фона ячейки"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_image_with_caption(doc, image_path, caption):
    """Добавить изображение с подписью"""
    if os.path.exists(image_path):
        # Центрирующий параграф для изображения
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(5.5))
        
        # Подпись под изображением
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_p.add_run(caption)
        caption_run.font.size = Pt(10)
        caption_run.font.italic = True
        caption_run.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()  # Пустая строка
    else:
        p = doc.add_paragraph(f"[Изображение не найдено: {image_path}]")
        p_run = p.runs[0]
        p_run.font.italic = True
        p_run.font.color.rgb = RGBColor(255, 0, 0)

def create_practice_1_docx():
    """Создать DOCX документ для Практической работы № 1"""
    
    doc = Document()
    
    # Установить стили шрифта
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ПРАКТИЧЕСКАЯ РАБОТА № 1")
    run.font.size = Pt(16)
    run.font.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Анализ логических выражений и функций")
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Дата
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("21 января 2026 г.")
    run.font.size = Pt(11)
    run.font.italic = True
    
    doc.add_paragraph()  # Пустая строка
    
    # ===== ЗАДАНИЕ 1 =====
    heading1 = doc.add_paragraph()
    run = heading1.add_run("1.1 Задание 1: Логическое выражение")
    run.font.size = Pt(12)
    run.font.bold = True
    
    # Условие задачи
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Дано логическое выражение:")
    run.font.italic = True
    run.font.size = Pt(11)
    
    formula1 = doc.add_paragraph()
    formula1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula1.add_run("ā·b̄·c̄ + ā·b·c + a·b̄·c + a·b·c̄ = y")
    run.font.size = Pt(12)
    run.font.bold = True
    
    # Диаграмма логического выражения
    add_image_with_caption(
        doc,
        r'e:\computer___technology\Pasted image 20250922125016.png',
        "Рисунок 1.1 - Диаграмма логического выражения"
    )
    
    doc.add_paragraph()
    
    # ===== ЗАДАНИЕ 2 =====
    heading2 = doc.add_paragraph()
    run = heading2.add_run("1.2 Задание 2: Анализ функции F")
    run.font.size = Pt(12)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p_run = p.add_run("Исходная функция:")
    p_run.font.bold = True
    p_run.font.size = Pt(11)
    
    formula2 = doc.add_paragraph()
    formula2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula2.add_run("F = ¯(Ā·B ∨ B·C) ∨ (Ā·C)")
    run.font.size = Pt(12)
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Таблица истинности с картинкой
    add_image_with_caption(
        doc,
        r'e:\computer___technology\Pasted image 20250922125035.png',
        "Рисунок 1.2 - Таблица истинности функции F"
    )
    
    # Таблица истинности текстовая
    p = doc.add_paragraph("Таблица истинности функции F:")
    run = p.runs[0]
    run.font.bold = True
    run.font.size = Pt(11)
    
    table = doc.add_table(rows=9, cols=10)
    table.style = 'Light Grid Accent 1'
    
    # Заголовок таблицы
    headers = ['A', 'B', 'C', 'Ā', 'X', 'Y', 'Z', 'U', 'V', 'F']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Подсветка заголовка
        set_cell_background(header_cells[i], 'D3D3D3')
        # Центрирование текста
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # Данные таблицы
    data = [
        ['0', '0', '0', '1', '0', '0', '0', '1', '0', '1'],
        ['0', '0', '1', '1', '0', '0', '0', '1', '1', '1'],
        ['0', '1', '0', '1', '1', '0', '1', '0', '0', '0'],
        ['0', '1', '1', '1', '1', '1', '1', '0', '1', '1'],
        ['1', '0', '0', '0', '0', '0', '0', '1', '0', '1'],
        ['1', '0', '1', '0', '0', '0', '0', '1', '0', '1'],
        ['1', '1', '0', '0', '0', '0', '0', '1', '0', '1'],
        ['1', '1', '1', '0', '0', '1', '1', '0', '0', '0'],
    ]
    
    for row_idx, row_data in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = cell_data
            # Центрирование текста
            for paragraph in row_cells[col_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            
            # Подсветка результата F (последний столбец)
            if col_idx == 9 and cell_data == '1':
                set_cell_background(row_cells[col_idx], 'FFFF99')
    
    doc.add_paragraph()
    
    # Вывод
    p = doc.add_paragraph()
    run = p.add_run("Вывод:")
    run.font.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph(
        "Функция F принимает значение 1 для всех наборов входных переменных, "
        "кроме (0,1,0) и (1,1,1)."
    )
    
    # Сохранение
    output_path = r'e:\computer___technology\Практическая 1.docx'
    doc.save(output_path)
    print(f"✓ DOCX файл создан: {output_path}")

if __name__ == '__main__':
    create_practice_1_docx()
